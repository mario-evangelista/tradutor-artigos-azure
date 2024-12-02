import requests
from docx import Document

# Configurações da API do Azure
AZURE_API_KEY = "SUA_CHAVE_AZURE"
AZURE_ENDPOINT = "SEU_ENDPOINT"
AZURE_REGION = "SUA_REGIAO"

# Função para traduzir texto
def traduzir_texto(texto, idioma_destino):
    url = f"{AZURE_ENDPOINT}/translate?api-version=3.0&to={idioma_destino}"
    headers = {
        "Ocp-Apim-Subscription-Key": AZURE_API_KEY,
        "Ocp-Apim-Subscription-Region": AZURE_REGION,
        "Content-Type": "application/json",
    }
    body = [{"text": texto}]
    
    response = requests.post(url, headers=headers, json=body)
    if response.status_code == 200:
        traducoes = response.json()
        return traducoes[0]["translations"][0]["text"]
    else:
        raise Exception(f"Erro: {response.status_code} - {response.text}")

# Função para carregar e traduzir um arquivo DOCX
def traduzir_arquivo_docx(caminho_entrada, caminho_saida, idioma_destino):
    print(f"Lendo o arquivo {caminho_entrada}...")
    doc = Document(caminho_entrada)
    doc_traduzido = Document()
    
    for paragrafo in doc.paragraphs:
        texto_original = paragrafo.text
        if texto_original.strip():  # Evita traduzir parágrafos vazios
            print(f"Traduzindo: {texto_original[:50]}...")
            texto_traduzido = traduzir_texto(texto_original, idioma_destino)
        else:
            texto_traduzido = ""
        doc_traduzido.add_paragraph(texto_traduzido)
    
    print(f"Salvando o arquivo traduzido em {caminho_saida}...")
    doc_traduzido.save(caminho_saida)
    print("Tradução concluída!")

# Exemplo de uso
if __name__ == "__main__":
    # Configurações do arquivo
    caminho_entrada = "docs/artigo_original.docx"
    caminho_saida = "docs/artigo_traduzido.docx"
    idioma_destino = "en"  # Código do idioma destino (ex: 'en' para inglês)
    
    try:
        traduzir_arquivo_docx(caminho_entrada, caminho_saida, idioma_destino)
    except Exception as e:
        print(f"Erro ao traduzir o arquivo: {e}")
